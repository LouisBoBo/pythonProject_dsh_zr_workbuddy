"""导入解析器：Swagger/OpenAPI 接口文档 & 数据字典文档（xlsx/csv/txt/docx）。

全部离线解析。PDF 无可用库时返回明确提示。
"""

import io
import re
from typing import Dict, List

import yaml

# ---------------- 通用 ----------------

TABLE_KEYS = ["work_orders", "equipment_events", "quality_records"]
TABLE_NAMES = {
    "work_orders": "生产工单",
    "equipment_events": "设备事件",
    "quality_records": "质检记录",
    "other": "其他",
}

# 关键词权重：2 = 强特征词（归属明确），1 = 弱特征词
TABLE_KEYWORDS = {
    "work_orders": [
        (2, ["工单", "订单", "报工", "work", "order", "wo", "mo"]),
        (1, ["计划", "产品", "产线", "班次", "数量", "plan", "product", "line", "mo"]),
    ],
    "equipment_events": [
        (2, ["设备", "机台", "停机", "故障", "维修", "equipment", "device", "machine", "downtime"]),
        (1, ["事件", "时长", "运行", "event", "duration"]),
    ],
    "quality_records": [
        (2, ["质检", "检验", "抽检", "不良", "缺陷", "quality", "inspect", "defect", "check"]),
        (1, ["合格", "质量"]),
    ],
}


def _score_table(t: str, key: str):
    s = 0
    strong = 0
    for weight, kws in TABLE_KEYWORDS[key]:
        for kw in kws:
            if kw in t:
                s += weight
                if weight == 2:
                    strong += 1
    return s, strong


def guess_table(text: str) -> str:
    """按加权关键词打分猜测归属数据表。"""
    t = (text or "").lower()
    scored = [(k, *_score_table(t, k)) for k in TABLE_KEYS]
    scored.sort(key=lambda x: (-x[1], -x[2]))
    best = scored[0]
    return best[0] if best[1] > 0 else "other"

# 常见字段名的标准字段自动建议
STD_SUGGEST = {
    "工单号": "work_order_no", "工单编号": "work_order_no", "订单号": "work_order_no",
    "产品编码": "product_code", "产品编号": "product_code", "物料编码": "product_code",
    "产品名称": "product_name", "物料名称": "product_name",
    "产线": "line_name", "线别": "line_name", "产线名称": "line_name",
    "班次": "shift", "班组": "shift",
    "计划数量": "planned_qty", "计划数": "planned_qty", "计划产量": "planned_qty",
    "实际数量": "actual_qty", "实际数": "actual_qty", "完成数量": "actual_qty", "产量": "actual_qty",
    "开始时间": "start_time", "开始日期": "start_time", "开工时间": "start_time",
    "结束时间": "end_time", "完成时间": "end_time", "结束日期": "end_time",
    "状态": "status", "工单状态": "status",
    "设备编码": "equipment_code", "设备编号": "equipment_code", "机台编码": "equipment_code",
    "设备名称": "equipment_name", "机台名称": "equipment_name",
    "事件类型": "event_type", "设备事件": "event_type", "类型": "event_type",
    "时长": "duration_min", "时长(分钟)": "duration_min", "持续时间": "duration_min",
    "检验时间": "inspect_time", "检验日期": "inspect_time",
    "抽检数": "sample_qty", "抽检数量": "sample_qty", "检验数量": "sample_qty",
    "不良数": "defect_qty", "不良数量": "defect_qty",
    "缺陷类型": "defect_type", "不良类型": "defect_type", "不良原因": "defect_type",
    # ERP 常见字段
    "order_no": "work_order_no", "plan_quantity": "planned_qty", "actual_quantity": "actual_qty",
    "plan_qty": "planned_qty", "actual_qty": "actual_qty",
    "production_line": "line_name", "good_count": "good_qty", "defect_count": "defect_qty",
    "scrap_count": "scrap_qty", "total_inspected": "sample_qty",
    "record_date": "record_date", "inspect_date": "inspect_date", "output_qty": "actual_qty",
    "defect_type": "defect_type", "quantity": "quantity", "status": "status",
}

TYPE_MAP = {
    "文本": "文本", "string": "文本", "varchar": "文本", "char": "文本", "str": "文本",
    "数值": "数值", "int": "数值", "integer": "数值", "number": "数值", "decimal": "数值",
    "float": "数值", "double": "数值", "bigint": "数值", "numeric": "数值",
    "时间": "时间", "datetime": "时间", "date": "时间", "timestamp": "时间", "time": "时间",
    "布尔": "布尔", "boolean": "布尔", "bool": "布尔",
}


def guess_table(text: str) -> str:
    """按加权关键词打分猜测归属数据表。"""
    t = (text or "").lower()
    scored = [(k, *_score_table(t, k)) for k in TABLE_KEYS]
    scored.sort(key=lambda x: (-x[1], -x[2]))
    best = scored[0]
    return best[0] if best[1] > 0 else "other"


def suggest_std(mes_field: str) -> str:
    f = (mes_field or "").strip()
    if f in STD_SUGGEST:
        return STD_SUGGEST[f]
    slug = re.sub(r"[^a-zA-Z0-9_\u4e00-\u9fff]", "_", f).strip("_")
    return slug


def norm_type(t: str) -> str:
    t = (t or "").strip().lower()
    t = re.sub(r"\(.*\)", "", t).strip()  # VARCHAR(50) → varchar
    return TYPE_MAP.get(t, "文本")


# ERP 等系统表名 → 标准表映射（从文档 `#### 表名` 标题识别）
TABLE_FROM_NAME = {
    "work_orders": "work_orders", "kanban_boards": "work_orders",
    "production_plans": "work_orders", "production_output_records": "work_orders",
    "wip_snapshots": "work_orders", "line_capacity_snapshots": "work_orders",
    "equipment": "equipment_events", "equipment_runtime_logs": "equipment_events",
    "equipment_oee_snapshots": "equipment_events", "equipment_alarms": "equipment_events",
    "equipment_output_records": "equipment_events",
    "equipment_maintenance_plans": "equipment_events", "equipment_maintenance_orders": "equipment_events",
    "equipment_repairs": "equipment_events", "equipment_repair_parts": "equipment_events",
    "quality_metrics": "quality_records", "quality_anomalies": "quality_records",
    "quality_defect_details": "quality_records",
}

# 通用噪音字段（跳过）
NOISE_FIELDS = {"id", "created_at", "updated_at", "remark"}


# ---------------- Swagger / OpenAPI ----------------

def parse_swagger(content: str) -> dict:
    data = yaml.safe_load(content)
    if not isinstance(data, dict) or "paths" not in data:
        raise ValueError("不是有效的 OpenAPI/Swagger 文档（缺少 paths 节点）")
    info = data.get("info") or {}
    endpoints = []
    for path, item in (data.get("paths") or {}).items():
        if not isinstance(item, dict):
            continue
        for method in ("get", "post", "put", "delete", "patch"):
            op = item.get(method)
            if not isinstance(op, dict):
                continue
            name = op.get("operationId") or op.get("summary") or f"{method.upper()} {path}"
            desc = op.get("summary") or ""
            params = []
            for p in op.get("parameters") or []:
                if isinstance(p, dict) and p.get("name"):
                    params.append(f"{p.get('name')}({p.get('in', '')})")
            if params and not desc:
                desc = "参数: " + ", ".join(params[:6])
            tags = " ".join(op.get("tags") or [])
            table = guess_table(f"{name} {path} {desc} {tags}")
            endpoints.append({
                "name": name, "method": method.upper(), "path": path,
                "table": table, "desc": desc,
            })
    if not endpoints:
        raise ValueError("Swagger 文档中未解析到任何接口（paths 为空？）")
    return {
        "endpoints": endpoints,
        "info": {"title": info.get("title") or "", "version": info.get("version") or "",
                 "count": len(endpoints)},
    }


# ---------------- Markdown 接口文档 ----------------

_METHODS = ("GET", "POST", "PUT", "DELETE", "PATCH")


def _is_sep_row(cells) -> bool:
    """Markdown 表格分隔行：|---|---| 等。"""
    if not cells:
        return False
    return all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells) and any("-" in c for c in cells)


def parse_markdown_api(content: str) -> dict:
    """解析 Markdown 接口文档（表格 / 请求方式+请求路径 / GET /path 三种写法）。"""
    endpoints = []
    seen = set()
    cur_heading = ""
    pending = None  # 等待路径的方法（"请求方式：GET" 与 "请求路径：/x" 分行时）
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^#{1,6}\s+(.*)$", line)
        if m:
            cur_heading = m.group(1).strip()
            continue

        # 1) 表格行
        if "|" in line:
            cells = [c.strip(" `*") for c in line.strip("|").split("|")]
            if _is_sep_row(cells):
                continue
            meth_idx = next((i for i, c in enumerate(cells)
                             if c.upper() in _METHODS or c.upper().startswith(_METHODS)), None)
            path_idx = next((i for i, c in enumerate(cells)
                             if i != meth_idx and (c.startswith("/") or re.match(r"^[A-Za-z][\w.\-]*/", c))), None)
            if meth_idx is None or path_idx is None:
                continue
            meth = re.match(r"(GET|POST|PUT|DELETE|PATCH)", cells[meth_idx], re.I).group(1).upper()
            path = cells[path_idx]
            name = cells[0] if (cells and meth_idx != 0 and path_idx != 0 and cells[0]) else (cur_heading or f"{meth} {path}")
            desc = " ".join(c for i, c in enumerate(cells)
                            if i not in (meth_idx, path_idx) and i != 0 and c)
            key = (meth, path)
            if key in seen:
                continue
            seen.add(key)
            endpoints.append({
                "name": name, "method": meth, "path": path,
                "table": guess_table(f"{name} {path} {desc} {cur_heading}"), "desc": desc,
            })
            continue

        # 2) "请求方式：GET" / "请求路径：/api/xxx"（可跨行）
        m = re.search(r"(?:请求方式|请求方法|请求类型|method)[:：]\s*(GET|POST|PUT|DELETE|PATCH)", line, re.I)
        if m:
            pending = m.group(1).upper()
            continue
        p = re.search(r"(?:请求路径|请求地址|接口地址|接口路径|path|url)[:：]\s*([^\s，,;；]+)", line, re.I)
        if p:
            path = p.group(1)
            meth = pending or next((mm for mm in _METHODS if f" {mm} " in f" {line} "), None)
            pending = None
            if meth:
                key = (meth, path)
                if key not in seen:
                    seen.add(key)
                    desc = line
                    name = cur_heading or f"{meth} {path}"
                    endpoints.append({
                        "name": name, "method": meth, "path": path,
                        "table": guess_table(f"{name} {path} {desc} {cur_heading}"),
                        "desc": desc[:80],
                    })
            continue

        # 3) 行内 "GET /api/v1/workOrders 说明"
        m = re.search(r"\b(GET|POST|PUT|DELETE|PATCH)\s+([/\w.\-{}:]+)", line, re.I)
        if m:
            meth, path = m.group(1).upper(), m.group(2)
            desc = line.replace(m.group(0), "").strip(" -–—:：|")
            key = (meth, path)
            if key in seen:
                continue
            seen.add(key)
            name = (cur_heading if any(k in cur_heading for k in ("接口", "查询", "列表", "获取", "新增", "删除", "修改", "更新"))
                    else desc or cur_heading or f"{meth} {path}")
            endpoints.append({
                "name": name, "method": meth, "path": path,
                "table": guess_table(f"{name} {path} {desc} {cur_heading}"), "desc": desc,
            })

    if not endpoints:
        raise ValueError("未能从 Markdown 中解析到接口（需要 方法+路径，如 GET /api/v1/workOrders，或接口表格）")
    return {
        "endpoints": endpoints,
        "info": {"title": "Markdown 接口文档", "version": "", "count": len(endpoints)},
    }


def parse_api_doc(filename: str, content: str) -> dict:
    """接口文档统一入口：优先 Swagger/OpenAPI，失败自动尝试 Markdown 解析。"""
    try:
        return parse_swagger(content)
    except Exception as e1:  # 含 yaml.ParserError 等（Markdown 常被 YAML 误解析）
        try:
            return parse_markdown_api(content)
        except ValueError as e2:
            raise ValueError(
                f"既不是有效的 Swagger/OpenAPI 文档（{e1}），也不是 Markdown 接口文档（{e2}）。"
                "支持：OpenAPI JSON/YAML，或 Markdown（接口表格 / 请求方式+请求路径 / GET /path）"
            ) from e2


# ---------------- 数据字典 ----------------

HEADER_ALIASES = {
    "mes_field": ["字段", "列名", "字段名", "字段名称", "名称", "参数", "属性", "field", "column", "param", "key"],
    "std_field": ["标准字段", "标准", "映射", "std", "standard"],
    "type": ["类型", "type", "数据类型", "格式", "字段类型"],
    "desc": ["说明", "描述", "含义", "备注", "desc", "description", "comment"],
}


def _norm(v) -> str:
    return "" if v is None else str(v).strip()


def _map_header(cols: List[str]) -> Dict[int, str]:
    roles = {}
    for i, c in enumerate(cols):
        cl = _norm(c).lower()
        for role, aliases in HEADER_ALIASES.items():
            if any(a.lower() in cl for a in aliases):
                roles[i] = role
                break
    return roles


def _role_index(roles: Dict[int, str], role: str):
    for i, r in roles.items():
        if r == role:
            return i
    return None


def _rows_to_mappings(rows: List[List[str]], skip_header: bool = True) -> List[dict]:
    """二维表格 → 映射行。skip_header=True 且检测不到表头时跳过首行。"""
    if not rows:
        return []
    roles = _map_header(rows[0])
    has_header = bool(roles)
    data_rows = rows if (has_header or not skip_header) else rows[1:]
    mes_idx = _role_index(roles, "mes_field")
    std_idx = _role_index(roles, "std_field")
    typ_idx = _role_index(roles, "type")
    desc_idx = _role_index(roles, "desc")
    out = []
    for row in data_rows:
        cells = [_norm(c) for c in row]
        if not any(cells):
            continue
        if has_header:
            mes = cells[mes_idx] if mes_idx is not None else (cells[0] if cells else "")
            std = cells[std_idx] if std_idx is not None else ""
            typ = cells[typ_idx] if typ_idx is not None else ""
            desc = cells[desc_idx] if desc_idx is not None else ""
        else:
            mes = cells[0] if cells else ""
            typ = cells[1] if len(cells) > 1 else ""
            desc = cells[2] if len(cells) > 2 else ""
            std = ""
        if not mes or mes.lower() in ("字段", "字段名", "列名", "名称", "field", "column"):
            continue
        out.append({
            "mes_field": mes,
            "std_field": std or suggest_std(mes),
            "type": norm_type(typ),
            "desc": desc,
        })
    return out


def _classify_row(row: dict) -> str:
    t = f"{row['mes_field']} {row['std_field']} {row['desc']}".lower()
    scored = [(k, *_score_table(t, k)) for k in TABLE_KEYS]
    scored.sort(key=lambda x: (-x[1], -x[2]))
    best = scored[0]
    return best[0] if best[1] > 0 else "uncategorized"


def _split_groups(mappings: List[dict]) -> Dict[str, List[dict]]:
    groups = {k: [] for k in TABLE_KEYS + ["uncategorized"]}
    seen = {k: set() for k in TABLE_KEYS + ["uncategorized"]}
    for m in mappings:
        forced = TABLE_FROM_NAME.get((m.get("table") or "").lower())
        g = forced or _classify_row(m)
        if m["mes_field"] in seen[g]:
            continue
        seen[g].add(m["mes_field"])
        groups[g].append({"mes_field": m["mes_field"], "std_field": m["std_field"],
                          "type": m["type"], "desc": m["desc"]})
    return groups


def parse_dictionary(filename: str, data: bytes) -> dict:
    ext = (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    if ext == "pdf":
        raise ValueError("PDF 解析需要 pdfplumber 库（当前环境无网络无法安装）。请导出为 Excel/CSV/TXT/Word 后上传。")
    if ext in ("xlsx", "xlsm"):
        mappings = _parse_xlsx(data)
    elif ext == "csv":
        mappings = _parse_csv(data)
    elif ext in ("txt", "md", "text"):
        mappings = _parse_markdown_txt(data)
    elif ext == "docx":
        mappings = _parse_docx(data)
    else:
        raise ValueError(f"暂不支持 {ext or '未知'} 格式，支持：xlsx / csv / txt / docx")

    groups = _split_groups(mappings)
    return {
        "groups": groups,
        "info": {"filename": filename, "total": len(mappings),
                 "counts": {k: len(v) for k, v in groups.items()}},
    }


def _parse_xlsx(data: bytes) -> List[dict]:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            rows.append([_norm(c) for c in row])
    return _rows_to_mappings(rows)


def _parse_csv(data: bytes) -> List[dict]:
    import pandas as pd
    raw = None
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            raw = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if raw is None:
        raw = data.decode("utf-8", errors="replace")
    df = pd.read_csv(io.StringIO(raw), dtype=str)
    rows = [list(df.columns)] + df.fillna("").values.tolist()
    return _rows_to_mappings(rows)


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_txt(data: bytes) -> List[dict]:
    """纯文本解析（无字段表头的简单文档）。管道表格先拆列，避免空单元格。"""
    raw = _decode(data)
    rows = []
    for line in raw.splitlines():
        line = line.strip()
        if not line or line.startswith(("#", "//", "--")):
            continue
        line = re.sub(r"[*`~]", "", line)
        if "|" in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if _is_sep_row(cells):
                continue
            rows.append(cells)
        else:
            rows.append(re.split(r"\s{1,}", re.sub(r"\s*[，,]\s*", " ", line)))
    return _rows_to_mappings(rows, skip_header=False)


def _is_field_header(cells) -> bool:
    """严格判断是否为字段表头：含字段名列（排除外键/父键/键）+ 类型或说明列。"""
    texts = [_norm(c).lower() for c in cells]
    mes_ok = False
    has_other = False
    for t in texts:
        if not t:
            continue
        if any(x in t for x in ("外键", "父", "键")):
            continue
        if any(a in t for a in ("字段", "列名", "名称", "field", "column", "param", "属性")):
            mes_ok = True
        if any(a in t for a in ("类型", "type", "说明", "描述", "备注", "desc", "comment", "含义", "默认")):
            has_other = True
    return mes_ok and has_other


def _parse_markdown_lines(raw: str) -> List[tuple]:
    """逐行解析 Markdown：跟踪 `表名` 标题，返回 (table_hint, cells) 列表。"""
    rows = []
    cur_table = ""
    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            m = re.match(r"^#{1,6}\s+`([^`]+)`", line)
            if m:
                cur_table = m.group(1).strip().lower()
            continue
        line = re.sub(r"[*`~]", "", line)
        if "|" in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if _is_sep_row(cells):
                continue
            rows.append((cur_table, cells))
        else:
            rows.append((cur_table, re.split(r"\s{1,}", re.sub(r"\s*[，,]\s*", " ", line))))
    return rows


def _rows_to_mappings_md(rows: List[tuple]) -> List[dict]:
    """Markdown 分段解析：只在字段表头激活后收集行，并携带表名提示。"""
    def cell(cells, idx):
        return cells[idx] if idx is not None and idx < len(cells) else ""

    out = []
    roles = None
    for table, cells in rows:
        if _is_field_header(cells):
            roles = _map_header(cells)
            continue
        if roles is None or not any(cells):
            continue
        mes_idx = _role_index(roles, "mes_field")
        std_idx = _role_index(roles, "std_field")
        typ_idx = _role_index(roles, "type")
        desc_idx = _role_index(roles, "desc")
        mes = cell(cells, mes_idx) or (cells[0] if cells else "")
        std = cell(cells, std_idx)
        typ = cell(cells, typ_idx)
        desc = cell(cells, desc_idx)
        if not mes:
            continue
        if mes.lower() in NOISE_FIELDS:
            continue
        if mes in ("—", "-") or re.fullmatch(r"\d+\.?", mes):
            continue
        if any(x in mes for x in ("：", "→", "**", "关系", "其余")):
            continue
        out.append({
            "mes_field": mes, "std_field": std or suggest_std(mes),
            "type": norm_type(typ), "desc": desc, "table": table,
        })
    return out


def _parse_markdown_txt(data: bytes) -> List[dict]:
    raw = _decode(data)
    rows = _parse_markdown_lines(raw)
    mappings = _rows_to_mappings_md(rows)
    if not mappings:
        # 兜底：无字段表头的简单文档（每行 字段 类型 说明）
        plain = [cells for _, cells in rows]
        mappings = _rows_to_mappings(plain, skip_header=False)
    return mappings


def _parse_docx(data: bytes) -> List[dict]:
    import docx
    doc = docx.Document(io.BytesIO(data))
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([_norm(c.text) for c in row.cells])
    if not rows:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        rows = [re.split(r"\s{1,}", re.sub(r"\s*[|，,]\s*", " ", l)) for l in lines]
    return _rows_to_mappings(rows)
